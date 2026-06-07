from ddgs import DDGS
import pymupdf
import pandas as pd
import os
import base64
from google import genai
from google.genai import types as genai_types
from bs4 import BeautifulSoup
import requests
import yt_dlp
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import sounddevice as sd
import soundfile as sf
import tempfile
from faster_whisper import WhisperModel


def save_to_xlsm(filename: str, data: list[dict]):
    """
    Saves structured data into an Excel Macro-Enabled Workbook (.xlsm).
    Args:
        filename: The name of the file (e.g., 'data.xlsm').
        data: A list of dictionaries representing rows.
    """
    try:
        df = pd.DataFrame(data)
        
        # Force the extension to be .xlsm if it isn't already
        if not filename.lower().endswith('.xlsm'):
            filename = os.path.splitext(filename)[0] + ".xlsm"
            
        # We use openpyxl as the engine to support the xlsm format
        df.to_excel(filename, index=False, engine='openpyxl')
        
        return f"Successfully saved {len(data)} rows to {filename} as a Macro-Enabled Workbook."
    except Exception as e:
        return f"Error saving .xlsm file: {str(e)}"


def file_manager(filename: str, mode: str, content: str = ""):
    """
    Reads, writes, or creates any text-based file.
    Args:
        filename: The name of the file (e.g., 'notes.txt', 'report.md').
        mode: What to do - 'read' to read, 'write' to overwrite, 'create' to make a new file.
        content: The text to write (only needed for 'write' and 'create' modes).
    Returns:
        str: File content (read mode) or a success/error message.
    """
    try:
        if mode == "read":
            if not os.path.exists(filename):
                return f"Error: '{filename}' does not exist."
            with open(filename, 'r') as f:
                return f.read()

        elif mode == "create":
            if os.path.exists(filename):
                return f"Error: '{filename}' already exists. Use mode='write' to overwrite it."
            with open(filename, 'w') as f:
                f.write(content)
            return f"'{filename}' created successfully."

        elif mode == "write":
            with open(filename, 'w') as f:
                f.write(content)
            return f"'{filename}' written successfully."

        else:
            return f"Error: Unknown mode '{mode}'. Use 'read', 'write', or 'create'."

    except Exception as e:
        return f"File operation failed: {str(e)}"



def internet_search(query: str):
    """
    Searches the internet for real-time information.
    Args:
        query: The search term (e.g. 'Bitcoin price today')
    Returns:
        str: Top 3 search results with titles, snippets and links.
    """
    try:
        results = DDGS().text(query, max_results=3)
        if not results:
            return "No results found."
        
        formatted = []
        for r in results:
            formatted.append(f"Title: {r['title']}\nSnippet: {r['body']}\nLink: {r['href']}")
        
        return "\n\n".join(formatted)
    
    except Exception as e:
        return f"Search failed: {str(e)}"
    



   

def read_pdf(filename: str, pages: str = "all"):
    """
    Reads and extracts text from a PDF file.
    Args:
        filename: Path to the PDF file (e.g., 'report.pdf').
        pages: Which pages to read - 'all' for entire PDF, or a page number like '1', '3'.
    Returns:
        str: Extracted text content from the PDF.
    """
    try:
        if not os.path.exists(filename):
            return f"Error: '{filename}' does not exist."

        if not filename.lower().endswith('.pdf'):
            return f"Error: '{filename}' is not a PDF file."

        doc = pymupdf.open(filename)
        total_pages = len(doc)

        if pages == "all":
            extracted = []
            for i, page in enumerate(doc):
                text = page.get_text().strip()
                if text:
                    extracted.append(f"--- Page {i + 1} ---\n{text}")
            doc.close()

            if not extracted:
                return "No readable text found. The PDF may be scanned/image-based."

            return f"PDF: '{filename}' ({total_pages} pages)\n\n" + "\n\n".join(extracted)

        else:
            page_num = int(pages) - 1  # Convert to 0-based index
            if page_num < 0 or page_num >= total_pages:
                return f"Error: Page {pages} doesn't exist. PDF has {total_pages} pages."
            text = doc[page_num].get_text().strip()
            doc.close()
            return f"PDF: '{filename}' — Page {pages}:\n\n{text}"

    except ValueError:
        return "Error: 'pages' must be 'all' or a valid page number like '1'."
    except Exception as e:
        return f"PDF reading failed: {str(e)}"
    


def read_image(filename: str, question: str = "Describe this image in detail."):
    """
    Reads and analyzes an image file using Gemini's vision capability.
    Args:
        filename: Path to the image file (e.g., 'photo.jpg', 'chart.png').
        question: What to ask about the image (e.g., 'Summarize the text in this image').
    Returns:
        str: Gemini's analysis/description of the image.
    """
    try:
        SUPPORTED = ['.jpg', '.jpeg', '.png', '.webp', '.gif', '.bmp']
        MIME_TYPES = {
            '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
            '.png': 'image/png',  '.webp': 'image/webp',
            '.gif': 'image/gif',  '.bmp': 'image/bmp'
        }

        if not os.path.exists(filename):
            return f"Error: '{filename}' does not exist."

        ext = os.path.splitext(filename)[1].lower()
        if ext not in SUPPORTED:
            return f"Error: Unsupported format '{ext}'. Supported: {SUPPORTED}"

        # Read and encode image as base64
        with open(filename, 'rb') as f:
            image_data = base64.standard_b64encode(f.read()).decode('utf-8')

        # Send to Gemini vision
        vision_client = genai.Client(api_key=os.environ.get("Gemini_API_Key"))
        response = vision_client.models.generate_content(
            model="gemini-2.5-flash-lite",
            contents=[
                genai_types.Part.from_bytes(
                    data=base64.b64decode(image_data),
                    mime_type=MIME_TYPES[ext]
                ),
                question
            ]
        )

        return f"Image analysis for '{filename}':\n\n{response.text}"

    except Exception as e:
        return f"Image reading failed: {str(e)}"
    


def scrape_webpage(url: str, extract: str = "text"):
    """
    Scrapes content from any webpage.
    Args:
        url: The full URL to scrape (e.g., 'https://example.com').
        extract: What to extract:
                 'text'   - all readable text (default)
                 'tables' - all tables as structured data
                 'links'  - all hyperlinks
                 'images' - all image URLs
    Returns:
        str: Extracted content from the webpage.
    """
    try:
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/120.0.0.0 Safari/537.36"
            )
        }

        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "lxml")

        # Remove junk tags
        for tag in soup(["script", "style", "nav", "footer", "ads"]):
            tag.decompose()

        # --- TEXT ---
        if extract == "text":
            text = soup.get_text(separator="\n")
            # Clean up blank lines
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            return "\n".join(lines)

        # --- TABLES ---
        elif extract == "tables":
            tables = soup.find_all("table")
            if not tables:
                return "No tables found on this page."

            all_tables = []
            for i, table in enumerate(tables):
                rows = []
                for tr in table.find_all("tr"):
                    cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                    if cells:
                        rows.append(" | ".join(cells))
                all_tables.append(f"Table {i+1}:\n" + "\n".join(rows))

            return "\n\n".join(all_tables)

        # --- LINKS ---
        elif extract == "links":
            links = []
            for a in soup.find_all("a", href=True):
                text = a.get_text(strip=True)
                href = a["href"]
                # Make relative URLs absolute
                if href.startswith("/"):
                    base = "/".join(url.split("/")[:3])
                    href = base + href
                if text:
                    links.append(f"{text} → {href}")
            return "\n".join(links) if links else "No links found."

        # --- IMAGES ---
        elif extract == "images":
            images = []
            for img in soup.find_all("img", src=True):
                src = img.get("src", "")
                alt = img.get("alt", "no description")
                if src.startswith("/"):
                    base = "/".join(url.split("/")[:3])
                    src = base + src
                images.append(f"Alt: {alt} → {src}")
            return "\n".join(images) if images else "No images found."

        else:
            return f"Unknown extract type '{extract}'. Use: 'text', 'tables', 'links', or 'images'."

    except requests.exceptions.ConnectionError:
        return f"Error: Could not connect to '{url}'. Check the URL."
    except requests.exceptions.Timeout:
        return "Error: Request timed out. The site may be slow."
    except requests.exceptions.HTTPError as e:
        return f"HTTP Error: {e}"
    except Exception as e:
        return f"Scraping failed: {str(e)}"



def download_youtube(url: str, mode: str = "video", resolution: str = "1080p", output_dir: str = "."):
    """
    Downloads a YouTube video or audio.
    Args:
        url: YouTube video URL.
        mode: 'video' (default) or 'audio'.
        resolution: '144p','240p','360p','480p','720p','1080p','1440p','4k'
                    Falls back to highest available if chosen resolution not found.
        output_dir: Folder to save file. Default is current folder.
    Returns:
        str: Success message or error.
    """
    try:
        os.makedirs(output_dir, exist_ok=True)

        target_height = int(
            resolution.lower()
            .replace('p', '')
            .replace('4k', '2160')
            .replace('k', '000')
        )

        if mode == "audio":
            ydl_opts = {
                "format": "bestaudio/best",
                "outtmpl": f"{output_dir}/%(title)s.%(ext)s",
                "quiet": True,
                "no_warnings": True,
            }

        elif mode == "video":
            ydl_opts = {
                # Download pre-merged formats only — no ffmpeg needed
                "format": (
                    f"best[height<={target_height}][ext=mp4]/"
                    f"best[height<={target_height}]/"
                    "best[ext=mp4]/best"
                ),
                "outtmpl": f"{output_dir}/%(title)s.%(ext)s",
                "quiet": True,
                "no_warnings": True,
            }

        else:
            return f"Error: Unknown mode '{mode}'. Use 'video' or 'audio'."

        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            title = info.get("title", "Unknown")
            duration = info.get("duration", 0)
            minutes = duration // 60
            seconds = duration % 60

            print(f"[Downloading] {title} ({minutes}m {seconds}s)...")
            ydl.download([url])

        return (
            f"✅ Download complete!\n"
            f"Title: {title}\n"
            f"Duration: {minutes}m {seconds}s\n"
            f"Mode: {mode.capitalize()}\n"
            f"Resolution: {resolution}\n"
            f"Saved to: {output_dir}/"
        )

    except yt_dlp.utils.DownloadError as e:
        return f"Download failed: {str(e)}"
    except ValueError:
        return f"Error: Invalid resolution '{resolution}'. Use: 144p, 240p, 360p, 480p, 720p, 1080p, 1440p, 4k"
    except Exception as e:
        return f"Unexpected error: {str(e)}"



def manage_docx(filename: str, mode: str, content: str = "", formatting: dict = {}):
    """
    Reads or writes Microsoft Word (.docx) files.
    Args:
        filename: Path to the file (e.g., 'report.docx').
        mode: What to do:
              'read'   - Read and return all text content.
              'create' - Create a new docx file with content.
              'write'  - Add new content to an existing docx file.
              'replace'- Overwrite entire content of existing file.
        content: Text to write. Use '\\n' for new lines, '---' for page break.
                 Prefix lines with:
                 '# '  for Title
                 '## ' for Heading 1
                 '### 'for Heading 2
                 '-'   for bullet points
        formatting: Optional dict for styling:
                    {'font_size': 12, 'bold': True, 'align': 'center'}
    Returns:
        str: File content (read mode) or success message.
    """
    try:
        # ── READ ──────────────────────────────────────────────
        if mode == "read":
            if not os.path.exists(filename):
                return f"Error: '{filename}' does not exist."
            doc = Document(filename)
            lines = []
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text)
            # Also read tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    lines.append(" | ".join(cells))
            return "\n".join(lines) if lines else "Document is empty."

        # ── CREATE ────────────────────────────────────────────
        elif mode == "create":
            if os.path.exists(filename):
                return f"Error: '{filename}' already exists. Use mode='replace' to overwrite."
            doc = Document()
            _write_content_to_doc(doc, content, formatting)
            doc.save(filename)
            return f"✅ '{filename}' created successfully."

        # ── WRITE (append) ────────────────────────────────────
        elif mode == "write":
            if not os.path.exists(filename):
                return f"Error: '{filename}' does not exist. Use mode='create' first."
            doc = Document(filename)
            _write_content_to_doc(doc, content, formatting)
            doc.save(filename)
            return f"✅ Content appended to '{filename}' successfully."

        # ── REPLACE (overwrite) ───────────────────────────────
        elif mode == "replace":
            doc = Document()
            _write_content_to_doc(doc, content, formatting)
            doc.save(filename)
            return f"✅ '{filename}' overwritten successfully."

        else:
            return f"Error: Unknown mode '{mode}'. Use: 'read', 'create', 'write', or 'replace'."

    except Exception as e:
        return f"Docx operation failed: {str(e)}"


def _write_content_to_doc(doc: Document, content: str, formatting: dict):
    """
    Internal helper — parses content string and writes
    formatted paragraphs, headings, and bullets to a Document.
    """
    font_size = formatting.get("font_size", 12)
    bold      = formatting.get("bold", False)
    align     = formatting.get("align", "left")

    alignment_map = {
        "left":    WD_ALIGN_PARAGRAPH.LEFT,
        "center":  WD_ALIGN_PARAGRAPH.CENTER,
        "right":   WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()  # blank line
            continue

        if line == "---":
            doc.add_page_break()

        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)   # Title

        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)   # Heading 1

        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)   # Heading 2

        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")

        else:
            para = doc.add_paragraph()
            para.alignment = alignment_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
            run = para.add_run(line)
            run.bold = bold
            run.font.size = Pt(font_size)


def speak(text: str, model: str = None):
    """
    Speaks text out loud using local Piper TTS.
    Args:
        text: The text to speak out loud.
        model: Optional path to a specific .onnx voice model.
    Returns:
        str: Confirmation message.
    """
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))

        # Correct path to piper executable
        piper_path = os.path.join(base_dir, "piper", "piper")

        # Correct path to voice model
        if not model:
            model = os.path.join(base_dir, "piper-voices", "en_US-ryan-high.onnx")

        # Check both exist before running
        if not os.path.exists(piper_path):
            return f"Error: Piper not found at '{piper_path}'"
        if not os.path.exists(model):
            return f"Error: Voice model not found at '{model}'"

        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
            temp_path = f.name

        process = subprocess.run(
            [piper_path, "--model", model, "--output_file", temp_path],
            input=text.encode("utf-8"),
            capture_output=True
        )

        if process.returncode != 0:
            return f"Piper error: {process.stderr.decode()}"

        data, sample_rate = sf.read(temp_path)
        sd.play(data, sample_rate)
        sd.wait()
        os.unlink(temp_path)

        return f"Spoken: '{text}'"

    except FileNotFoundError:
        return "Error: 'piper' executable not found."
    except Exception as e:
        return f"Speak failed: {str(e)}"
    

# Loads once when read_files.py is imported
print("[Voice] Loading Whisper model...")
whisper_model = WhisperModel("base", device="cpu", compute_type="int8")
print("[Voice] Whisper ready.")

def listen(duration: int = 6):
    """
    Records your voice from the microphone and converts it to text.
    Args:
        duration: How many seconds to record (default 6).
    Returns:
        str: Transcribed text from your speech.
    """
    try:
        sample_rate = 16000

        print(f"[Listening for {duration} seconds... speak now!]")
        audio = sd.rec(
            int(duration * sample_rate),
            samplerate=sample_rate,
            channels=1,
            dtype="float32"
        )
        sd.wait()  # Wait until recording finishes
        print("[Processing your speech...]")

        # Save to temp file for Whisper
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
            temp_path = f.name
        sf.write(temp_path, audio, sample_rate)

        # Transcribe
        segments, _ = whisper_model.transcribe(temp_path, beam_size=5)
        text = " ".join(seg.text for seg in segments).strip()

        # Clean up temp file
        os.unlink(temp_path)

        return text if text else ""

    except Exception as e:
        return f"Listen failed: {str(e)}"